import csv
import io
from datetime import date, datetime, time, timedelta, timezone
from io import BytesIO
from uuid import UUID

from docx import Document
from openpyxl import Workbook
from fastapi import HTTPException, status
from sqlalchemy import and_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.timezone import MOSCOW_TZ, add_seconds_to_start, format_moscow_date, format_moscow_datetime
from app.models.stream import BroadcastSession, SponsorMention, StreamEvent
from app.schemas.report import ReportMentionRow, ReportMentionsOut
from app.utils.timecode import seconds_to_hhmmss


def _range_utc_moscow_days(date_from: date, date_to: date) -> tuple[datetime, datetime]:
    start_local = datetime.combine(date_from, time.min, tzinfo=MOSCOW_TZ)
    end_exclusive_local = datetime.combine(date_to + timedelta(days=1), time.min, tzinfo=MOSCOW_TZ)
    return start_local.astimezone(timezone.utc), end_exclusive_local.astimezone(timezone.utc)


async def get_mentions_report(
    session: AsyncSession,
    *,
    stream_event_id: UUID | None,
    date_from: date | None,
    date_to: date | None,
) -> ReportMentionsOut:
    q = (
        select(SponsorMention)
        .join(BroadcastSession)
        .join(StreamEvent)
        .options(
            selectinload(SponsorMention.broadcast_session).selectinload(BroadcastSession.stream_event),
        )
    )
    conds = []
    if stream_event_id is not None:
        conds.append(StreamEvent.id == stream_event_id)
    if date_from is not None and date_to is not None:
        start_utc, end_exc = _range_utc_moscow_days(date_from, date_to)
        conds.append(and_(SponsorMention.created_at >= start_utc, SponsorMention.created_at < end_exc))
    elif date_from is not None or date_to is not None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Укажите обе даты диапазона (date_from и date_to)",
        )
    if conds:
        q = q.where(and_(*conds))
    # Сначала турнир (дата старта, id), затем день эфира, затем порядок упоминаний — без перемешивания турниров
    q = q.order_by(
        StreamEvent.start_date.asc(),
        StreamEvent.id.asc(),
        BroadcastSession.day_index.asc(),
        BroadcastSession.started_at.asc(),
        SponsorMention.created_at.asc(),
    )
    result = await session.execute(q)
    rows: list[ReportMentionRow] = []
    for m in result.scalars().all():
        bs = m.broadcast_session
        ev = bs.stream_event
        started = bs.started_at if bs.started_at.tzinfo else bs.started_at.replace(tzinfo=timezone.utc)
        abs_adj = add_seconds_to_start(started, m.adjusted_offset_sec)
        event_day_date = ev.start_date + timedelta(days=bs.day_index - 1)
        rows.append(
            ReportMentionRow(
                mention_id=m.id,
                stream_event_id=ev.id,
                stream_title=ev.title,
                event_day_date=event_day_date,
                day_index=bs.day_index,
                broadcast_session_id=bs.id,
                original_timecode=seconds_to_hhmmss(m.original_offset_sec),
                adjusted_timecode=seconds_to_hhmmss(m.adjusted_offset_sec),
                absolute_moscow_adjusted=format_moscow_datetime(abs_adj),
                is_adjusted=m.original_offset_sec != m.adjusted_offset_sec,
                mention_created_at=m.created_at,
            )
        )
    return ReportMentionsOut(items=rows, total=len(rows))


def build_docx_report(rows: list[ReportMentionRow]) -> bytes:
    doc = Document()
    rows_sorted = sorted(
        rows,
        key=lambda r: (r.stream_event_id, r.day_index, r.mention_created_at),
    )
    current_key: tuple[UUID, int] | None = None
    mention_idx = 0
    for row in rows_sorted:
        key = (row.stream_event_id, row.day_index)
        if key != current_key:
            if current_key is not None:
                doc.add_paragraph("")
            current_key = key
            mention_idx = 0
            doc.add_heading(row.stream_title, level=1)
            doc.add_paragraph(f"Дата: {format_moscow_date(row.event_day_date)}")
            doc.add_paragraph(f"День эфира: {row.day_index}")
        mention_idx += 1
        doc.add_paragraph(
            f"Упоминание {mention_idx} — таймкод {row.adjusted_timecode}, "
            f"абсолютное (МСК): {row.absolute_moscow_adjusted}, "
            f"запись: {format_moscow_datetime(row.mention_created_at)}",
        )
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_csv_report(rows: list[ReportMentionRow]) -> bytes:
    rows_sorted = sorted(rows, key=lambda r: (r.stream_event_id, r.day_index, r.mention_created_at))
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(
        [
            "stream_title",
            "event_day_date",
            "day_index",
            "original_timecode",
            "adjusted_timecode",
            "absolute_moscow_adjusted",
            "mention_created_at",
        ]
    )
    for row in rows_sorted:
        w.writerow(
            [
                row.stream_title,
                row.event_day_date.isoformat(),
                row.day_index,
                row.original_timecode,
                row.adjusted_timecode,
                row.absolute_moscow_adjusted,
                row.mention_created_at.isoformat(),
            ]
        )
    return buf.getvalue().encode("utf-8-sig")


def build_xlsx_report(rows: list[ReportMentionRow]) -> bytes:
    rows_sorted = sorted(rows, key=lambda r: (r.stream_event_id, r.day_index, r.mention_created_at))
    wb = Workbook()
    ws = wb.active
    ws.title = "mentions"
    ws.append(
        [
            "stream_title",
            "event_day_date",
            "day_index",
            "original_timecode",
            "adjusted_timecode",
            "absolute_moscow_adjusted",
            "mention_created_at",
        ]
    )
    for row in rows_sorted:
        ws.append(
            [
                row.stream_title,
                row.event_day_date.isoformat(),
                row.day_index,
                row.original_timecode,
                row.adjusted_timecode,
                row.absolute_moscow_adjusted,
                row.mention_created_at.isoformat(),
            ]
        )
    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


async def export_mentions_docx(
    session: AsyncSession,
    *,
    stream_event_id: UUID | None,
    date_from: date | None,
    date_to: date | None,
) -> bytes:
    data = await get_mentions_report(
        session,
        stream_event_id=stream_event_id,
        date_from=date_from,
        date_to=date_to,
    )
    return build_docx_report(data.items)


async def export_mentions_csv(
    session: AsyncSession,
    *,
    stream_event_id: UUID | None,
    date_from: date | None,
    date_to: date | None,
) -> bytes:
    data = await get_mentions_report(
        session,
        stream_event_id=stream_event_id,
        date_from=date_from,
        date_to=date_to,
    )
    return build_csv_report(data.items)


async def export_mentions_xlsx(
    session: AsyncSession,
    *,
    stream_event_id: UUID | None,
    date_from: date | None,
    date_to: date | None,
) -> bytes:
    data = await get_mentions_report(
        session,
        stream_event_id=stream_event_id,
        date_from=date_from,
        date_to=date_to,
    )
    return build_xlsx_report(data.items)
